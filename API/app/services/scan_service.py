import os
import io
import json
from typing import Dict, Any, List, Optional
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
import openpyxl
from app.utils.common import normalize_path, handle_exceptions, logger

@handle_exceptions
def read_docx(path: str) -> Dict[str, Any]:
    """
    Đọc và phân tích file Word (.docx)
    """
    norm_path = normalize_path(path)
    
    if not os.path.exists(norm_path):
        raise FileNotFoundError(f"File does not exist: {path}")
    
    if not norm_path.lower().endswith('.docx'):
        raise ValueError(f"File is not a Word document: {path}")
    
    # Đọc file docx
    doc = Document(norm_path)
    
    # Lấy nội dung text
    paragraphs = [p.text for p in doc.paragraphs]
    
    # Lấy bảng
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    return {
        "path": norm_path,
        "size": os.path.getsize(norm_path),
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": "\n".join(paragraphs)
    }

@handle_exceptions
def read_excel(path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
    """
    Đọc và phân tích file Excel (.xlsx, .xls)
    """
    norm_path = normalize_path(path)
    
    if not os.path.exists(norm_path):
        raise FileNotFoundError(f"File does not exist: {path}")
    
    if not any(norm_path.lower().endswith(ext) for ext in ['.xlsx', '.xls']):
        raise ValueError(f"File is not an Excel document: {path}")
    
    # Đọc file excel
    workbook = openpyxl.load_workbook(norm_path, read_only=True, data_only=True)
    
    # Lấy danh sách sheet
    sheet_names = workbook.sheetnames
    
    # Nếu không chỉ định sheet_name, lấy sheet đầu tiên
    if sheet_name is None:
        sheet_name = sheet_names[0]
    elif sheet_name not in sheet_names:
        raise ValueError(f"Sheet '{sheet_name}' not found in workbook")
    
    # Lấy dữ liệu từ sheet
    sheet = workbook[sheet_name]
    
    # Chuyển đổi sang dataframe
    data = []
    for row in sheet.iter_rows(values_only=True):
        data.append(list(row))
    
    df = pd.DataFrame(data[1:], columns=data[0] if data else None)
    
    return {
        "path": norm_path,
        "size": os.path.getsize(norm_path),
        "sheet_names": sheet_names,
        "current_sheet": sheet_name,
        "data": df.to_dict(orient="records"),
        "shape": df.shape
    }

@handle_exceptions
def read_pdf(path: str) -> Dict[str, Any]:
    """
    Đọc và phân tích file PDF
    """
    norm_path = normalize_path(path)
    
    if not os.path.exists(norm_path):
        raise FileNotFoundError(f"File does not exist: {path}")
    
    if not norm_path.lower().endswith('.pdf'):
        raise ValueError(f"File is not a PDF document: {path}")
    
    # Đọc file PDF
    pdf = PdfReader(norm_path)
    
    # Lấy thông tin cơ bản
    info = pdf.metadata
    num_pages = len(pdf.pages)
    
    # Lấy text từ từng trang
    pages_text = []
    for i in range(num_pages):
        page = pdf.pages[i]
        pages_text.append(page.extract_text())
    
    return {
        "path": norm_path,
        "size": os.path.getsize(norm_path),
        "info": {
            "title": info.title,
            "author": info.author,
            "subject": info.subject,
            "creator": info.creator,
            "producer": info.producer
        },
        "num_pages": num_pages,
        "pages": pages_text,
        "full_text": "\n".join(pages_text)
    }

@handle_exceptions
def read_csv(path: str, delimiter: str = ",") -> Dict[str, Any]:
    """
    Đọc và phân tích file CSV
    """
    norm_path = normalize_path(path)
    
    if not os.path.exists(norm_path):
        raise FileNotFoundError(f"File does not exist: {path}")
    
    # Đọc file CSV với pandas
    df = pd.read_csv(norm_path, delimiter=delimiter)
    
    return {
        "path": norm_path,
        "size": os.path.getsize(norm_path),
        "columns": df.columns.tolist(),
        "data": df.to_dict(orient="records"),
        "shape": df.shape
    }

@handle_exceptions
def read_text_file(path: str) -> Dict[str, Any]:
    """
    Đọc file text thông thường (.txt, .md, .py, .java, etc.)
    """
    norm_path = normalize_path(path)
    
    if not os.path.exists(norm_path):
        raise FileNotFoundError(f"File does not exist: {path}")
    
    # Thử đọc với encoding utf-8
    try:
        with open(norm_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # Nếu không đọc được với utf-8, thử với latin-1
        with open(norm_path, 'r', encoding='latin-1') as f:
            content = f.read()
    
    # Lấy số dòng
    lines = content.split('\n')
    
    return {
        "path": norm_path,
        "size": os.path.getsize(norm_path),
        "content": content,
        "line_count": len(lines)
    }

@handle_exceptions
def read_file_auto(path: str) -> Dict[str, Any]:
    """
    Tự động nhận diện loại file và đọc với phương thức thích hợp
    """
    norm_path = normalize_path(path)
    
    if not os.path.exists(norm_path):
        raise FileNotFoundError(f"File does not exist: {path}")
    
    ext = os.path.splitext(norm_path)[1].lower()
    
    # Phân loại file theo extension
    if ext == '.docx':
        return read_docx(norm_path)
    elif ext in ['.xlsx', '.xls']:
        return read_excel(norm_path)
    elif ext == '.pdf':
        return read_pdf(norm_path)
    elif ext == '.csv':
        return read_csv(norm_path)
    elif ext in ['.txt', '.md', '.py', '.java', '.html', '.css', '.js', '.json', '.xml']:
        return read_text_file(norm_path)
    else:
        # Cho các loại file chưa được hỗ trợ, đọc thử như plain text
        try:
            return read_text_file(norm_path)
        except UnicodeDecodeError:
            raise ValueError(f"Unsupported file type or cannot decode file: {path}") 